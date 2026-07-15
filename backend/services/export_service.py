"""
Export service — generate PDF, DOCX, and TXT exports of meeting data.
"""

import io
import logging
from datetime import datetime

from fastapi.responses import StreamingResponse

logger = logging.getLogger(__name__)


class ExportService:
    """Export meeting data in various formats."""

    def export_txt(self, meeting, transcript, summary, action_items) -> StreamingResponse:
        """Export meeting as plain text."""
        lines = []
        lines.append("=" * 60)
        lines.append(f"MEETING SUMMARY: {meeting.title}")
        lines.append("=" * 60)
        lines.append("")

        if meeting.meeting_date:
            lines.append(f"Date: {meeting.meeting_date.strftime('%Y-%m-%d %H:%M')}")
        if meeting.duration_seconds:
            minutes = meeting.duration_seconds // 60
            lines.append(f"Duration: {minutes} minutes")
        lines.append(f"Status: {meeting.status}")
        lines.append("")

        # Summary
        if summary:
            lines.append("-" * 40)
            lines.append("EXECUTIVE SUMMARY")
            lines.append("-" * 40)
            lines.append(summary.executive_summary)
            lines.append("")

            if summary.key_points:
                lines.append("KEY POINTS:")
                for i, point in enumerate(summary.key_points, 1):
                    lines.append(f"  {i}. {point}")
                lines.append("")

            if summary.decisions:
                lines.append("DECISIONS:")
                for i, decision in enumerate(summary.decisions, 1):
                    lines.append(f"  {i}. {decision}")
                lines.append("")

        # Action Items
        if action_items:
            lines.append("-" * 40)
            lines.append("ACTION ITEMS")
            lines.append("-" * 40)
            for i, item in enumerate(action_items, 1):
                lines.append(f"  {i}. [{item.priority.upper()}] {item.description}")
                if item.assignee:
                    lines.append(f"     Assignee: {item.assignee}")
                if item.due_date:
                    lines.append(f"     Due: {item.due_date}")
                lines.append(f"     Status: {item.status}")
            lines.append("")

        # Transcript
        if transcript:
            lines.append("-" * 40)
            lines.append("FULL TRANSCRIPT")
            lines.append("-" * 40)
            lines.append(transcript.full_text)
            lines.append("")

        lines.append("=" * 60)
        lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        lines.append("AI Meeting Summarizer")

        content = "\n".join(lines)
        buffer = io.BytesIO(content.encode("utf-8"))

        return StreamingResponse(
            buffer,
            media_type="text/plain",
            headers={
                "Content-Disposition": f'attachment; filename="{meeting.title}_summary.txt"'
            },
        )

    def export_docx(self, meeting, transcript, summary, action_items) -> StreamingResponse:
        """Export meeting as DOCX."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(meeting.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if meeting.meeting_date:
            doc.add_paragraph(
                f"Date: {meeting.meeting_date.strftime('%Y-%m-%d %H:%M')}"
            )
        if meeting.duration_seconds:
            doc.add_paragraph(f"Duration: {meeting.duration_seconds // 60} minutes")

        # Summary
        if summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(summary.executive_summary)

            if summary.key_points:
                doc.add_heading("Key Points", level=2)
                for point in summary.key_points:
                    doc.add_paragraph(point, style="List Bullet")

            if summary.decisions:
                doc.add_heading("Decisions", level=2)
                for decision in summary.decisions:
                    doc.add_paragraph(decision, style="List Bullet")

        # Action Items
        if action_items:
            doc.add_heading("Action Items", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            headers = table.rows[0].cells
            headers[0].text = "Description"
            headers[1].text = "Assignee"
            headers[2].text = "Priority"
            headers[3].text = "Status"

            for item in action_items:
                row = table.add_row().cells
                row[0].text = item.description
                row[1].text = item.assignee or "Not Mentioned"
                row[2].text = item.priority
                row[3].text = item.status

        # Transcript
        if transcript:
            doc.add_heading("Full Transcript", level=1)
            doc.add_paragraph(transcript.full_text)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} | AI Meeting Summarizer"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{meeting.title}_summary.docx"'
            },
        )

    def export_pdf(self, meeting, transcript, summary, action_items) -> StreamingResponse:
        """Export meeting as PDF using ReportLab."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75 * inch)
        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=20,
            textColor=colors.HexColor("#1a1a2e"),
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading1"],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.HexColor("#16213e"),
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=10,
            spaceAfter=6,
            leading=14,
        )

        # Title
        story.append(Paragraph(meeting.title, title_style))

        # Metadata
        meta_parts = []
        if meeting.meeting_date:
            meta_parts.append(f"Date: {meeting.meeting_date.strftime('%Y-%m-%d %H:%M')}")
        if meeting.duration_seconds:
            meta_parts.append(f"Duration: {meeting.duration_seconds // 60} min")
        if meta_parts:
            story.append(Paragraph(" | ".join(meta_parts), body_style))
        story.append(Spacer(1, 12))

        # Summary
        if summary:
            story.append(Paragraph("Executive Summary", heading_style))
            story.append(Paragraph(summary.executive_summary, body_style))
            story.append(Spacer(1, 8))

            if summary.key_points:
                story.append(Paragraph("Key Points", heading_style))
                for point in summary.key_points:
                    story.append(Paragraph(f"• {point}", body_style))
                story.append(Spacer(1, 8))

            if summary.decisions:
                story.append(Paragraph("Decisions", heading_style))
                for decision in summary.decisions:
                    story.append(Paragraph(f"• {decision}", body_style))
                story.append(Spacer(1, 8))

        # Action Items
        if action_items:
            story.append(Paragraph("Action Items", heading_style))
            table_data = [["#", "Description", "Assignee", "Priority", "Status"]]
            for i, item in enumerate(action_items, 1):
                table_data.append([
                    str(i),
                    item.description[:80],
                    item.assignee or "N/A",
                    item.priority,
                    item.status,
                ])

            table = Table(table_data, colWidths=[30, 200, 80, 60, 60])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f0f0")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

        # Transcript
        if transcript:
            story.append(Paragraph("Full Transcript", heading_style))
            # Split long transcript into chunks for PDF rendering
            text = transcript.full_text
            chunk_size = 2000
            for i in range(0, len(text), chunk_size):
                chunk = text[i : i + chunk_size].replace("\n", "<br/>")
                story.append(Paragraph(chunk, body_style))

        # Footer
        story.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            "Footer", parent=body_style, fontSize=8, textColor=colors.grey
        )
        story.append(
            Paragraph(
                f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} | AI Meeting Summarizer",
                footer_style,
            )
        )

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{meeting.title}_summary.pdf"'
            },
        )


# Singleton
export_service = ExportService()
